from django.shortcuts import render
from django.http import HttpResponse
from django import forms
from docx import Document
from django.conf import settings
from django.templatetags.static import static
import os
import zipfile
from io import BytesIO
from decimal import Decimal, InvalidOperation
from django.shortcuts import get_object_or_404
from administrateur.models import ProformaInscrit,ProformaNonInscrit, DetailClient
from num2words import num2words


# Vue pour charger les informations du proforma
def generer_rie(request, proforma_id):
    SECTEURS_CHOICES = [
        ('Textile habillement et Accessoire (THA)', 'Textile habillement et Accessoire (THA)'),
        ('Développement rural (DR)', 'Développement rural (DR)'),
        ('BTP-Ressources Stratégiques (BTP/RS)', 'BTP-Ressources Stratégiques (BTP/RS)'),
        ('Tourisme Hôtellerie Restauration (THR)', 'Tourisme Hôtellerie Restauration (THR)'),
        ('TIC', 'TIC'),
        ('Autres', 'Autres'),
    ]

    try:
        # Vérifie d'abord si le proforma est inscrit
        proforma = ProformaInscrit.objects.get(id=proforma_id)
        client = get_object_or_404(DetailClient, utilisateur=proforma.idclient)
        initial_data = {
            'raison_sociale': client.raison_social if client.type_client == 2 else f"{client.utilisateur.nom} {client.utilisateur.prenom}",
            'cnaps': client.num_cnaps or '',
            'adresse': client.adresse or '',
            'effectif': client.effectif or 0,
            'email': client.utilisateur.email,
            'contact': client.contact_entreprise or '',
            'cout_total': proforma.cout_total
        }
    except ProformaInscrit.DoesNotExist:
        # Sinon, le proforma est non inscrit
        proforma = get_object_or_404(ProformaNonInscrit, id=proforma_id)
        initial_data = {
            'raison_sociale': proforma.nom_entreprise,
            'cnaps': proforma.cnaps or '',
            'effectif': '',
            'email': '',
            'adresse': proforma.adresse_entreprise or '',
            'contact': proforma.contact or '',
            'cout_total': proforma.cout_total
        }

    return render(request, 'admin/fmfp/formulaire_rie.html', {'proforma': initial_data, 'secteur_choices': SECTEURS_CHOICES})


def remplir_formulaire(request):
    if request.method == 'POST':
        # Récupération des données soumises
        raison_sociale = request.POST.get('raison_sociale', '')
        cnaps = request.POST.get('cnaps', '')
        adresse = request.POST.get('adresse', '')
        effectif = request.POST.get('effectif', '')
        email = request.POST.get('email', '')
        contact = request.POST.get('contact', '')
        description_projet = request.POST.get('description_projet', '')
        secteurs = request.POST.getlist('secteurs')
        autre_secteur = request.POST.get('autre_secteur', '') if 'Autres' in secteurs else ''
        secteur_principal = ', '.join(secteurs)
        cout_total_str = request.POST.get('cout_total', '0')
        cout_total_str = cout_total_str.replace(',', '.')
        # Gestion sécurisée de cout_total
        try:
            cout_total = Decimal(cout_total_str)
        except InvalidOperation:
            cout_total = Decimal('0')

        # Séparer la partie entière et les centimes
        partie_entiere = int(cout_total)
        partie_decimale = int(round((cout_total - partie_entiere) * 100))

        # Conversion en lettres
        partie_entiere_lettres = num2words(partie_entiere, lang='fr')
        if partie_decimale > 0:
            partie_decimale_lettres = num2words(partie_decimale, lang='fr')
            cout_total_lettres = f"{partie_entiere_lettres} ariary et {partie_decimale_lettres} centimes"
        else:
            cout_total_lettres = f"{partie_entiere_lettres} ariary"

        # Chemins vers les documents à remplir
        rie_doc_path = os.path.join(settings.BASE_DIR, 'administrateur/static/fmfp/rie/Formulaire_RIE.docx')
        lettre_doc_path = os.path.join(settings.BASE_DIR, 'administrateur/static/fmfp/rie/Annexe-4_REI_Lettre_demande_financement.docx')

        # Dictionnaire des remplacements
        remplacements = {
            'raison_sociale': raison_sociale,
            'cnaps': cnaps,
            'adresse': adresse,
            'effectif': effectif,
            'email': email,
            'contact': contact,
            'cout_total': f"{cout_total} MGA",
            'cout_total_lettres': cout_total_lettres,
            'autre_secteur': autre_secteur,
            'description_projet': description_projet,
            'secteur': secteur_principal
        }

        # Fonction pour remplir les documents Word
        def remplir_document(doc_path, remplacements, secteurs):
            doc = Document(doc_path)
            for paragraphe in doc.paragraphs:
                remplacer_et_cocher(paragraphe, remplacements, secteurs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraphe in cell.paragraphs:
                            remplacer_et_cocher(paragraphe, remplacements, secteurs)
            return doc

        # Fonction pour remplacer les balises et cocher les cases
        def remplacer_et_cocher(paragraphe, remplacements, secteurs):
            for run in paragraphe.runs:
                for cle, valeur in remplacements.items():
                    if f'[{cle}]' in run.text:
                        run.text = run.text.replace(f'[{cle}]', str(valeur))
                for secteur in secteurs:
                    if secteur == 'Autres' and '☐ Autres' in run.text:
                        run.text = run.text.replace('☐ Autres', '☑ Autres')
                    elif f'☐ {secteur}' in run.text:
                        run.text = run.text.replace(f'☐ {secteur}', f'☑ {secteur}')

            if '[description_projet]' in paragraphe.text:
                paragraphe.text = paragraphe.text.replace('[description_projet]', '')
                paragraphe.add_run(description_projet)

        # Remplir les deux documents
        formulaire_rie = remplir_document(rie_doc_path, remplacements, secteurs)
        lettre_financement = remplir_document(lettre_doc_path, remplacements, secteurs)

        # Créer un fichier ZIP pour les deux documents
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
            rie_buffer = BytesIO()
            formulaire_rie.save(rie_buffer)
            zip_file.writestr('Formulaire_RIE_rempli.docx', rie_buffer.getvalue())

            lettre_buffer = BytesIO()
            lettre_financement.save(lettre_buffer)
            zip_file.writestr('Lettre_demande_financement_remplie.docx', lettre_buffer.getvalue())

        zip_buffer.seek(0)
        response = HttpResponse(zip_buffer, content_type='application/zip')
        response['Content-Disposition'] = 'attachment; filename=\"Documents_Formulaire_RIE.zip\"'
        return response

    return HttpResponse("Méthode non autorisée", status=405)